import os
from spire.doc import *
# from spire.doc.documents import Paragraph
from spire.doc import FileFormat, Document as SpireDoc  
from main.utilities.argos import translate
from docx import Document as docxDoc
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langdetect import detect
from pdf2docx import Converter
from docx import Document as docxDoc
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from main.models import *


def pdf2word(pdf_path, output_path):
    cv = Converter(pdf_path)
    cv.convert(output_path)
    cv.close()


def set_rtl(paragraph):
    p_pr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), 'true')
    p_pr.append(bidi)


def adjust_rtl(input_path, output_path):
    doc = docxDoc(input_path)
    mystyle_rtl = doc.styles.add_style('mystyle_rtl', WD_STYLE_TYPE.CHARACTER)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:  # Check if the paragraph is not empty
            try:
                for run in paragraph.runs:
                    run.style = mystyle_rtl
                    run.font.ltr = False
                    run.font.rtl = True
                set_rtl(paragraph)
            except Exception as e:
                print(f"Error processing paragraph: {e}")

    # Remove the first paragraph if added by Spire.Doc
    if doc.paragraphs[0].text.startswith(""):
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    doc.save(output_path)


def adjust_ltr(input_path, output_path):
    doc = docxDoc(input_path)
    
    # Create a paragraph style for LTR alignment
    paragraph_style = doc.styles.add_style('LTRParagraphStyle', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Create a character style for LTR text direction
    character_style = doc.styles.add_style('LTRCharacterStyle', WD_STYLE_TYPE.CHARACTER)
    character_style.font.rtl = False  # Ensure the text direction is LTR
    
    for paragraph in doc.paragraphs:
        # Apply the paragraph style
        paragraph.style = paragraph_style
        
        # Remove RTL markers if they exist
        p_pr = paragraph._element.get_or_add_pPr()
        rtl_element = p_pr.find(qn('w:bidi'))
        if rtl_element is not None:
            p_pr.remove(rtl_element)
        
        # Apply character style to runs
        for run in paragraph.runs:
            run.style = character_style
    
    # Save the document with adjusted formatting
    doc.save(output_path)


def doc_translator(input_path, output_path, source_lang, target_lang, translation_object_id):
    # Load the Word document
    translation_object = FileTranslationTask.objects.get(task_id=translation_object_id)

    document = SpireDoc()
    document.LoadFromFile(input_path)

    # Step 1: Extract images and save their positions
    image_positions = []
    images = []
    for i in range(len(document.Sections)):
        section = document.Sections[i]
        for j in range(len(section.Paragraphs)):
            paragraph = section.Paragraphs[j]
            for k in range(len(paragraph.ChildObjects)):
                child = paragraph.ChildObjects[k]
                if isinstance(child, DocumentObject):
                    if child.DocumentObjectType == DocumentObjectType.Picture:
                        images.append(child)
                        image_positions.append((i, j, k))  # Save position (section, paragraph, index)

    n_sections = len(document.Sections)
    # Step 2: Translate text and modify document
    for i in range(n_sections):
        print(f"Translating section {i + 1}/{n_sections}")
        section = document.Sections[i]
        for j in range(len(section.Paragraphs)):
            paragraph = section.Paragraphs[j]
            # Translate the paragraph text
            translated_text = translate(paragraph.Text, source_lang, target_lang)
            paragraph.Text = translated_text
            # Set the text direction to RTL for Persian
            if target_lang in ['fa', 'ar', 'he']:
                paragraph.Format.HorizontalAlignment = HorizontalAlignment.Right
            else:
                paragraph.Format.HorizontalAlignment = HorizontalAlignment.Left
            paragraph.Format.HorizontalAlignment = HorizontalAlignment.Justify
        translation_object.translated_percentage = int((i + 1) * 100 / n_sections)
        translation_object.save()

    # Step 3: Reinsert images at their original positions
    for idx, (sec_idx, para_idx, child_idx) in enumerate(image_positions):
        section = document.Sections[sec_idx]
        paragraph = section.Paragraphs[para_idx]
        paragraph.ChildObjects.Insert(child_idx, images[idx])

    # Save the translated and RTL formatted document
    document.SaveToFile(output_path, FileFormat.Docx2019)

    # Optional: Adjust for RTL-specific formatting if needed
    if target_lang in ['fa', 'ar', 'he']:
        adjust_rtl(output_path, output_path)
    else:
        adjust_ltr(output_path, output_path)

    print(f"Translation complete. The document has been saved to {output_path}.")


def file_translation_handler(input_path, output_path, source_lang, target_lang, translation_object_id):
    if input_path.endswith("pdf"):
        doc_path = os.path.splitext(input_path)[0] + ".docx"
        pdf2word(input_path, doc_path)
        # os.remove(input_path)
        input_path = doc_path
    elif input_path.endswith("txt"):
        pass # It should be completed later 
    elif os.path.splitext(input_path)[1] not in [".docx", ".doc", ".txt", ".pdf"]:
        raise TypeError("file's format is not supported. try: .docx, .doc, .pdf, .txt")

    doc_translator(input_path, output_path, source_lang, target_lang, translation_object_id)
    # os.remove(input_path)
